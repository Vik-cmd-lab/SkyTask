import os
import sys
import json
import threading
import webbrowser
from io import BytesIO
from docx import Document
from flask import Flask, render_template, request, redirect, url_for, send_file
from platformdirs import user_data_dir

APP_NAME = "SkyTask"
DATA_DIR = user_data_dir(APP_NAME)
DATA_FILE: str = os.path.join(DATA_DIR, "tasks.json")
os.makedirs(DATA_DIR, exist_ok=True)

print(DATA_DIR)

def load_tasks():
    global tasks, next_id
    tasks = []
    next_id = 1

    if not os.path.exists(DATA_FILE):
        save_tasks()
        return

    try:
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        tasks = data.get("tasks", [])
        next_id = data.get("next_id", 1)

    except (json.JSONDecodeError, OSError, AttributeError, TypeError):
        tasks = []
        next_id = 1
        save_tasks()


def save_tasks():
    global next_id
    if not tasks:
        next_id = 1

    temp_file = DATA_FILE + ".tmp"
    backup_file = DATA_FILE + ".bak"

    data = {
        "tasks": tasks,
        "next_id": next_id
    }

    try:
        with open(temp_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        if os.path.exists(DATA_FILE):
            os.replace(DATA_FILE, backup_file)

        os.replace(temp_file, DATA_FILE)

    except Exception:
        if os.path.exists(temp_file):
            os.remove(temp_file)
        raise

def get_path(relative_path: str) -> str:
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(str(sys._MEIPASS), relative_path)
    return os.path.abspath(relative_path)

app = Flask(
    __name__,
    template_folder=get_path("templates"),
    static_folder=get_path("static")
)

tasks = []
next_id: int = 1

bg_color = "#ffffff"
bg_image_url = ""

load_tasks()


@app.route('/')
def index():
    tasks.sort(key=lambda task: task["priority"])
    return render_template(
        "index.html",
        tasks=tasks,
        bg_color=bg_color,
        bg_image_url=bg_image_url,
        data_file=DATA_FILE,
        tasks_quantity=len(tasks)
    )

def open_browser():
    webbrowser.open('http://127.0.0.1:5000')

@app.route("/add", methods=['POST'])
def add_task():
    global next_id
    title = request.form.get('title')
    description = request.form.get('description')
    priority = request.form.get('priority')
    deadline = request.form.get('deadline')

    task = {
        'id': next_id,
        'title': title,
        'description': description,
        'priority': priority,
        'deadline': deadline,
        'done': False
    }

    tasks.append(task)
    next_id += 1
    save_tasks()

    return redirect(url_for('index'))


@app.route('/delete/<int:task_id>', methods=['POST'])
def delete_task(task_id):
    global tasks
    tasks = [task for task in tasks if task["id"] != task_id]
    save_tasks()
    return redirect(url_for('index'))


@app.route('/toggle/<int:task_id>', methods=['POST'])
def toggle_task(task_id):
    global tasks
    for task in tasks:
        if task["id"] == task_id:
            task["done"] = not task["done"]
            break
    save_tasks()
    return redirect(url_for('index'))


@app.route('/edit/<int:task_id>', methods=['POST'])
def edit_task(task_id):
    global tasks
    for task in tasks:
        if task["id"] == task_id:
            return render_template('edit.html', task=task)
    return redirect(url_for('index'))


@app.route('/update/<int:task_id>', methods=['POST'])
def update_task(task_id):
    global tasks
    for task in tasks:
        if task["id"] == task_id:
            task["title"] = request.form.get("title")
            task["description"] = request.form.get("description")
            task["priority"] = request.form.get("priority")
            task["deadline"] = request.form.get("deadline")
            break

    save_tasks()
    return redirect(url_for('index'))



@app.route("/background", methods=["POST"])
def set_background():
    global bg_color, bg_image_url
    bg_color = request.form.get("bg_color")
    bg_image_url = request.form.get("bg_image_url")
    return redirect(url_for("index"))


@app.route("/clear", methods=['POST'])
def clear_tasks():
    tasks.clear()
    save_tasks()
    return redirect(url_for('index'))


@app.route("/export", methods=["POST"])
def export_tasks():
    doc = Document()
    doc.add_heading("list of tasks", level=1)

    if not tasks:
        doc.add_paragraph("no tasks.")
    else:
        for task in tasks:
            p = doc.add_paragraph()
            p.add_run(f"ID: {task['id']}\n").bold = True
            p.add_run(f"Title: {task['title']}\n")
            p.add_run(f"Description: {task['description']}\n")
            p.add_run(f"Priority: {task['priority']}\n")
            p.add_run(f"Deadline: {task['deadline']}\n")
            p.add_run(f"Done: {task['done']}\n")
            doc.add_paragraph("-" * 20)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="tasks.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/stop", methods=["POST"])
def stop():
    os._exit(0)

if __name__ == "__main__":
    threading.Timer(1.0, open_browser).start()
    app.run(debug=False)